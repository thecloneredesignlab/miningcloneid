import argparse
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

OUT = Path(__file__).resolve().parents[1]


def set_cell_margins(*args, **kwargs):
    pass


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'PAGE'
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)


def configure(doc, title):
    sec = doc.sections[0]
    sec.top_margin = Inches(0.75)
    sec.bottom_margin = Inches(0.75)
    sec.left_margin = Inches(0.85)
    sec.right_margin = Inches(0.85)
    sec.header_distance = Inches(0.3)
    sec.footer_distance = Inches(0.35)

    styles = doc.styles
    normal = styles['Normal']
    normal.font.name = 'Arial'
    normal._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.08

    for name, size, bold in [('Title', 16, True), ('Heading 1', 13, True), ('Heading 2', 11.5, True), ('Heading 3', 10.5, True)]:
        st = styles[name]
        st.font.name = 'Arial'
        st._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
        st.font.size = Pt(size)
        st.font.bold = bold
        st.paragraph_format.keep_with_next = True
        st.paragraph_format.space_before = Pt(9 if name != 'Title' else 0)
        st.paragraph_format.space_after = Pt(4)

    footer = sec.footer.paragraphs[0]
    add_page_number(footer)

    p = doc.add_paragraph(style='Title')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(title)


def add_meta(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.italic = True
    r.font.size = Pt(9.5)


def add_label_paragraph(doc, label, text, mono=False):
    p = doc.add_paragraph()
    p.paragraph_format.keep_together = False
    r = p.add_run(label)
    r.bold = True
    r2 = p.add_run(text)
    if mono:
        r2.font.name = 'Courier New'
        r2._element.rPr.rFonts.set(qn('w:eastAsia'), 'Courier New')
        r2.font.size = Pt(9)
    return p


def add_numbered_comment(doc, num, title, body):
    h = doc.add_paragraph(style='Heading 2')
    h.add_run(f'{num}. {title}')
    for para in body:
        p = doc.add_paragraph(para)
        p.paragraph_format.keep_together = False


def add_minor(doc, num, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.15)
    p.paragraph_format.first_line_indent = Inches(-0.15)
    p.add_run(f'{num}. ').bold = True
    p.add_run(text)


def build_report():
    doc = Document()
    configure(doc, 'Peer Review Report')
    add_meta(doc, 'Manuscript: Resource limitation rewires chromosome instability and ploidy evolution across in vitro and in vivo cancer models')
    add_meta(doc, 'Recommendation: Major Revision')

    doc.add_heading('Overall Assessment', level=1)
    doc.add_paragraph(
        'This manuscript addresses an important and mechanistically interesting question: how resource context changes the production and selection of chromosome-number variation in matched near-diploid and near-tetraploid cancer lineages. The integration of passage-structured in vitro data, longitudinal in vivo tumor burden, terminal single-cell chromosome-number distributions, histologic necrosis, and a joint soft-coupled model is a substantial strength. Several findings are directionally stable across the selected solutions, particularly the low direct-death requirement for the severe-deprivation 4N chromosome-number decline, the lower effective proliferative ceiling in vivo, and the stronger nonzero-stress missegregation response in vivo.'
    )
    doc.add_paragraph(
        'However, the current manuscript cannot yet support its strongest mechanistic claims. The joint method is described incorrectly in the repository draft, the post-missegregation survival direction is reversed when the complete nonlinear function is evaluated, the joint results are strongly conditioned on an asymmetric warm-start design, and most fixed-oxygen response curves do not have strong spectral separation. In addition, the objective used in the Figure 6 ranking diagnostic is not the fitted MAP objective, the necrosis export is broken, and the provenance does not yet identify a reproducible clean code state. These are correctable issues, but they require substantive revision of the methods, figures, results, and interpretation.'
    )

    doc.add_heading('Major Comments', level=1)

    major_comments = [
        (
            'Correct the joint-model specification and make the manuscript match the implemented soft coupling',
            [
                'The repository draft states that alpha_o2 and gamma_growth are hard-shared and presents a quadratic Gaussian soft-coupling penalty for the remaining parameters. The current configuration, code path, README, and supplied joint outputs instead show that all 14 active biological parameters are represented by context-specific transformed center-delta values and are regularized by a bounded Welsch penalty. The manuscript must use one internally consistent specification throughout the main Methods, Supplementary Methods, Results, Figure 5 caption, and parameter descriptions.',
                'The Welsch form is not a cosmetic implementation detail. With c = 0.4, the penalty cap is 0.08 per parameter, and the selected solutions use approximately 75.8-89.5% of the total 14-parameter cap. Several headline splits are near saturation in all selected solutions. The authors should report the standardized deltas and saturation status and should avoid language implying that the estimated separations persisted under an indefinitely increasing quadratic shrinkage penalty.'
            ]
        ),
        (
            'Correct the altered-daughter survival interpretation and regenerate the relevant figure and text',
            [
                'The manuscript currently concludes that post-missegregation survival is more ploidy-dependent in vivo. Direct evaluation of the complete fitted survival function gives the opposite gradient. Across all 60 selected joint solutions, median per-copy survival is approximately 0.807 at 44 chromosomes and 0.911 at 88 chromosomes in vivo, compared with 0.204 and 0.837 in vitro. Thus, absolute altered-daughter survival is higher in vivo at both reference states, whereas the 44-to-88 chromosome survival gradient is much stronger in vitro (median absolute gradient 0.633 versus 0.096).',
                'This correction must be applied consistently to the Abstract, Results, Discussion, Figure 5D, Figure 5 caption, and any ratio panel. The safest presentation is to plot the derived survival functions, or function-level contrasts at biologically defined chromosome states, rather than to infer direction from s_max, beta_buf, and n_exp separately.'
            ]
        ),
        (
            'Demonstrate joint-optimization robustness and broaden the warm-start design',
            [
                'The six joint warm starts are objective-minimum representatives of six exploratory in vivo subclusters, but all six are paired with the same in vitro seed 10 from one 376-seed subcluster. Other in vitro subclusters are not represented. The six pair-specific best objectives range from 18.852 to 19.978, and all 60 selected local refinements were attempted but not accepted. Therefore, the selected joint fits are conditional local candidate solutions rather than a demonstrated global optimum or a balanced comparison of the two environments.',
                'The authors should expand the start design to include representative solutions from the in vitro landscape, or provide a systematic sensitivity analysis showing that the headline derived functions remain stable across a broader in vivo-by-in vitro start matrix. The convergence code, local-refinement rejection, stopping criteria, and movement from each warm start should be diagnosed explicitly. Until then, precise fold changes should be presented as ranges across selected local solutions, not as uniquely identified effects.'
            ]
        ),
        (
            'Quantify practical identifiability and bound sensitivity',
            [
                'The separate in vivo top-ten solutions have nearly equivalent objectives but widely different values for multiple mechanistic parameters, and several parameters reach calibration bounds. Sigma_burden is at its upper bound in all ten selected in vivo fits. The joint solutions also contain multiple boundary-limited context values and saturated soft-coupling dimensions. These patterns indicate practical non-identifiability or strong parameter compensation.',
                'A formal uncertainty analysis is required for the parameters and, more importantly, for the derived biological response functions used in the conclusions. Profile likelihood, parametric bootstrap, structured multistart envelopes, or an equivalent likelihood-based approach would be appropriate. The analysis should report whether the derived proliferation, missegregation, and survival contrasts remain bounded and directionally stable when nuisance parameters are re-optimized. Raue et al. specifically distinguish practical non-identifiability caused by limited data from structural non-identifiability and show how profile likelihood can identify flat directions and likelihood-based intervals (Bioinformatics 2009; Sections 2.2, 2.3, and 6; doi:10.1093/bioinformatics/btp358).'
            ]
        ),
        (
            'Integrate the full fixed-oxygen spectral-gap audit into the main interpretation',
            [
                'The newly supplied preprocessing permits direct validation of the full 500-seed fixed-oxygen analysis. The analysis contains 100,500 seed-oxygen attractor evaluations. Only 141 of 500 fits (28.2%) meet the recorded reliable spectral-gap criterion; 168 are cautionary and 191 are unreliable. None of the six in vivo warm-start representatives used for joint fitting is reliable: three are cautionary and three are unreliable.',
                'The manuscript should distinguish raw response shape from reliability in all fixed-oxygen figures and text. Small spectral gaps imply slow or ambiguous dominance between leading modes, so a curve label alone should not be interpreted as a stable long-term biological regime. The authors should show the reliability composition of every response class, identify the reliability of the selected warm starts, and perform sensitivity analyses on the gap thresholds or finite-time horizon.'
            ]
        ),
        (
            'Correct the objective definition used in the Figure 6 selection diagnostic',
            [
                'The fixed-oxygen pipeline selects objective_source = raw_likelihood for all 500 seeds when the automatic selector is used. This value is the raw burden-plus-ploidy likelihood diagnostic, not the complete fitted MAP objective that includes the normalized component structure, necrosis, and prior contribution. Its Spearman correlation with the separate-fit MAP objective is only approximately 0.429. Consequently, rings or rankings described as the best fitted or lowest-objective MAP solutions are mislabeled.',
                'The authors should regenerate the panel using an explicitly reconstructed MAP objective, or relabel it throughout as a raw-likelihood diagnostic and avoid using it to select a preferred MAP response class. The seed-level objective crosswalk used to generate the panel should be released. The exact Figure 4 AUC input and result tables should also be provided because their ranking could not be independently reproduced from the supplied archives.'
            ]
        ),
        (
            'Repair the necrosis observation export and align the equation, code, and report',
            [
                'The necrosis term is active in the fitted objective, but every inspected independent and joint necrosis_fit.tsv contains missing predicted values. Predicted necrosis can be reconstructed from the harvest-time dead-to-total burden ratio, and the reconstructed loss matches the reported objective only when the squared standardized logit residual is averaged without the factor of one-half shown in the repository draft equation. This reveals two distinct issues: a broken post-fit export and a factor-of-two Methods mismatch.',
                'The authors should repair the export so that observed value, predicted value, residual, standardized residual, inclusion status, and objective contribution are all written for each mapped sample. An automated test should reconstruct the necrosis objective from the exported table and compare it with fit_summary.tsv. The manuscript equation, code, and report must use exactly the same normalization.'
            ]
        ),
        (
            'Moderate the in vitro trajectory and individual-tumor fit claims',
            [
                'The severe-deprivation 4N result is supported: across the selected in vitro top-ten fits, mean chromosome number declines by 3.54-4.17 while direct hypoxia-origin dead burden remains 1.74-1.82%. In contrast, the 2N deprived model generates a mixed low/high chromosome distribution but does not reproduce the late divergence between the two experimental branches. The two observed final means are approximately 66.85 and 88.05, whereas the selected predictions are shared and approximately 63.72-64.40. The manuscript should not describe this as a uniformly transient high-ploidy expansion followed by an overall return to low ploidy.',
                'Similarly, the in vivo model gives the same prediction to tumors sharing starting cohort and harvest time because no mouse-specific latent effect or harvest-specific initial multiplier is fitted. The burden and endpoint summaries are useful at the cohort level, but the manuscript should not imply accurate individualized trajectories. Report the absolute fit-quality metrics and state explicitly which heterogeneity is represented versus left in residuals.'
            ]
        ),
        (
            'Separate biological nonviability from finite-grid boundary loss',
            [
                'The CIN-associated dead compartment receives both post-missegregation nonviable daughter mass and daughter or WGD mass routed outside the retained chromosome grid of 22-154. The in vitro selected solutions show a much larger CIN-associated fraction than direct hypoxia-origin death, but this fraction cannot be interpreted entirely as biological lethality while numerical boundary routing is included.',
                'The authors should export and plot missegregation-nonviable and boundary-dropped components separately and repeat the principal conclusions with expanded N_MIN/N_MAX values. A grid-sensitivity analysis is essential before using the magnitude of the CIN-associated compartment as a biological readout.'
            ]
        ),
        (
            'Provide a reproducible, figure-level release with clean provenance',
            [
                'The selected independent in vivo and joint runs record different commits and dirty working trees, while the selected in vitro directories do not contain equivalent run_provenance.tsv files. The full fixed-oxygen and clustering tables are now available, but the Figure 4 seed-level AUC tables remain absent. The current package therefore cannot identify one exact code state and complete intermediate-data chain for every figure.',
                'The final release should bind each result to a clean tagged commit or an archived patch, resolved configuration and parameter tables, exact input checksums, random seeds, software/container versions, intermediate tables, and a figure-to-source manifest. Sandve et al. emphasize recording how every result was produced, versioning custom scripts, retaining detailed output underlying summaries, and making scripts/runs/results accessible (PLOS Computational Biology 2013; Rules 1, 4, 8, and 10; doi:10.1371/journal.pcbi.1003285). The FAIR Principles additionally require detailed provenance for reusable digital research objects, including algorithms and workflows (Wilkinson et al., Scientific Data 2016; Box 2, R1.2; doi:10.1038/sdata.2016.18).'
            ]
        ),
    ]

    for i, (title, body) in enumerate(major_comments, 1):
        add_numbered_comment(doc, i, title, body)

    doc.add_heading('Minor Comments', level=1)
    minors = [
        'State the number of tumors with mapped necrosis measurements in the main Methods, not only in supplementary text.',
        'Use “chromosome number” when the modeled state is total chromosome count and reserve “ploidy” for normalized or categorical quantities. Define every conversion explicitly.',
        'Report oxygen units consistently as percent O2 and distinguish assigned in vitro oxygen from latent effective in vivo oxygen.',
        'Do not use “significant” for differences that have not undergone a defined statistical test. Use “directionally consistent across selected solutions” where appropriate.',
        'Describe top-ten and top-sixty ranges as optimizer-robustness summaries, not confidence intervals.',
        'Clarify that the WGD branch is implemented as an effective tracked-lineage conversion with unit retained mass, not as two viable daughters.',
        'Report all fixed-oxygen classification thresholds and the exact finite-difference rule version in a permanent supplementary table.',
        'For t-SNE, report perplexity, learning rate, iteration count, random seed, and stability across repeated embeddings. t-SNE cluster sizes and inter-cluster distances should not be given biological meaning; Wattenberg et al. discuss these specific interpretation failures in Sections 1-3 (Distill 2016; doi:10.23915/distill.00002).',
        'Reconcile panel lettering in Figure 6 and ensure that every panel mentioned in the Results exists in the assembled figure and has a source table.',
        'Remove all remaining editorial placeholders, unsupported exact causal language, and references to analyses whose input tables are not released.',
        'Add a Data and Code Availability section that names the permanent repository release, archive DOI, container or environment specification, and figure-generation entry point.',
        'Check all parameter symbols and subscripts for consistency between the main text, equations, parameter tables, code names, and plot labels.'
    ]
    for i, text in enumerate(minors, 1):
        add_minor(doc, i, text)

    doc.add_heading('Recommendation', level=1)
    doc.add_paragraph(
        'Major Revision. The central framework is potentially valuable, and several direction-level conclusions are supported. Publication readiness depends on correcting the joint method and survival direction, demonstrating broader optimization and identifiability robustness, integrating the spectral-gap audit, repairing the objective/export inconsistencies, and providing complete figure-level provenance.'
    )

    doc.add_heading('Methodological References Cited in This Review', level=1)
    refs = [
        'Raue A, Kreutz C, Maiwald T, et al. Structural and practical identifiability analysis of partially observed dynamical models by exploiting the profile likelihood. Bioinformatics. 2009;25:1923-1929. doi:10.1093/bioinformatics/btp358.',
        'Wattenberg M, Viégas F, Johnson I. How to Use t-SNE Effectively. Distill. 2016. doi:10.23915/distill.00002.',
        'Sandve GK, Nekrutenko A, Taylor J, Hovig E. Ten Simple Rules for Reproducible Computational Research. PLoS Comput Biol. 2013;9:e1003285. doi:10.1371/journal.pcbi.1003285.',
        'Wilkinson MD, Dumontier M, Aalbersberg IJJ, et al. The FAIR Guiding Principles for scientific data management and stewardship. Sci Data. 2016;3:160018. doi:10.1038/sdata.2016.18.'
    ]
    for ref in refs:
        doc.add_paragraph(ref)

    path = OUT / '05_peer_review_report.docx'
    doc.save(path)
    return path


def build_crosswalk():
    doc = Document()
    configure(doc, 'Major Comment Evidence Crosswalk')
    add_meta(doc, 'Detailed trace from manuscript statements to code, result files, verification, impact, and acceptance criteria')

    doc.add_heading('Purpose and Evidence Standard', level=1)
    doc.add_paragraph(
        'This document traces every major review comment to the repository manuscript, supplied result archives, repository implementation, independent recalculation, and any external methodological reference. “Verified” means that the stated quantity or implementation was directly reproduced from the supplied files or code. “Partially verified” means that the broad analysis is present but one or more source tables required for exact numerical reproduction are missing. No comment below is based solely on plausibility.'
    )

    entries = [
        {
            'title': 'Joint specification does not match the implemented 14-parameter Welsch soft coupling',
            'review': 'Rewrite the joint Methods so that all 14 active biological parameters are soft-coupled through transformed center-delta values and the implemented bounded Welsch penalty. Remove the hard-shared alpha_o2/gamma_growth description and the quadratic penalty equation.',
            'manuscript': 'Repository Supplementary Methods, “Joint fitting of in vivo and in vitro data,” states that alpha_o2 and gamma_growth are hard-shared, lists 12 soft-coupled parameters, includes an unresolved editorial note, and defines a quadratic penalty. Main Results and Figure 5 imply that all 14 active parameters are context-specific.',
            'evidence': 'Repo config `oxygen/config/O2_supply_demand.yaml` lists 14 names in `joint_soft_coupling_params`; Repo README and joint backend define center + delta/2 for in vivo, center - delta/2 for in vitro, union transformed bounds, and a Welsch penalty. Every supplied `joint_soft_coupling.tsv` contains the 14 context-specific parameters; the active hard-shared biological table is empty. The Welsch cap is 0.08 per parameter at c=0.4. `analysis/joint_soft_coupling_parameter_summary.csv` and `analysis/joint_selected60_fit_summary.csv` show substantial saturation.',
            'why': 'The parameterization determines the dimensionality, admissible domain, regularization geometry, and biological meaning of every context ratio. A bounded robust penalty behaves differently from a Gaussian prior: once saturated, additional separation is nearly unpenalized.',
            'external': 'No external theory is needed to establish the mismatch; it is a direct code/result audit.',
            'verification': 'Verified from config, README/backend semantics, and 60 selected joint outputs.',
            'impact': 'Critical. The current Methods are factually inconsistent with the fitted model and could lead readers to infer nonexistent hard constraints and stronger shrinkage than was used.',
            'acceptance': 'All manuscript sections and figures use the 14-parameter center-delta Welsch formulation; the exact formula, sigma, c, cap, bounds, and feasibility rules are stated; saturation diagnostics are reported.'
        },
        {
            'title': 'The fitted survival-function direction is reversed in the repository narrative',
            'review': 'Replace the claim of stronger ploidy-dependent survival in vivo with the function-level result: absolute altered-daughter survival is higher in vivo, but ploidy dependence is weaker in vivo and stronger in vitro.',
            'manuscript': 'Repository Results, “Tumors and culture impose distinct ploidy-selective pressures,” states that the post-missegregation survival function is more ploidy-dependent in vivo. The Discussion repeats this interpretation through lower s_max, lower beta_buf, and higher n_exp. Figure 5D/caption is framed in the same direction.',
            'evidence': 'The implemented function is `s_N = s_max * exp[-beta_buf*(44/N)^n_exp]`. Direct evaluation for every selected joint solution gives median in vivo s44=0.807 and s88=0.911, versus in vitro s44=0.204 and s88=0.837. Median s88/s44 is 1.12 in vivo and 4.11 in vitro; median absolute gradient is 0.096 versus 0.633. In 60/60 solutions, absolute survival is higher in vivo at both N values; in 0/60 is the in vivo gradient stronger. See `analysis/joint_survival_function_all60.csv` and `analysis/joint_survival_function_summary_all60.csv`.',
            'why': 'Nonlinear parameter combinations cannot be interpreted reliably from component directions. The biological prediction is the evaluated function, not the sign of an individual parameter change.',
            'external': 'No external theory is required; the correction follows by direct substitution into the manuscript’s own equation.',
            'verification': 'Verified across all 60 selected joint solutions.',
            'impact': 'Critical scientific direction error. It changes which environment is inferred to impose stronger ploidy discrimination after missegregation.',
            'acceptance': 'Abstract, Results, Discussion, Figure 5D, all captions, and summary statements display the corrected function-level direction and include the range across selected solutions.'
        },
        {
            'title': 'Joint solutions are conditioned on asymmetric exploratory warm starts and unaccepted local refinement',
            'review': 'Broaden the in vitro warm-start coverage, diagnose local-refinement failure, and present current joint results as local candidate regimes rather than a unique optimum.',
            'manuscript': 'Repository Results and Figure 5A describe six in vivo landscape-informed warm starts paired with the same best in vitro fit but do not explain the unrepresented in vitro subclusters, subcluster quality, pair-level objective spread, or local-refinement status.',
            'evidence': '`joint_pre.zip` shows six in vivo objective-minimum representatives from six subclusters and a single in vitro seed 10 from vt_C02_Sc02 (376/500 fits). Other in vitro subclusters are not represented. Pair-best objectives are 18.852, 18.890, 18.970, 19.414, 19.791, and 19.978. All 60 selected fits record local attempted but local accepted = false. See `analysis/joint_pair_summary.csv`, `analysis/joint_selected60_fit_summary.csv`, and warm-start movement tables.',
            'why': 'A one-sided start panel cannot establish robustness to culture-parameter uncertainty. Pair-level objective separation and failed local acceptance indicate that optimizer basin and bounds materially affect the selected solution.',
            'external': 'No external citation is necessary to establish the design limitation. It is an observed property of the supplied start matrix and optimizer outputs.',
            'verification': 'Verified from `joint_pre.zip` representative crosswalks and 60 joint seed outputs.',
            'impact': 'High. Precise context ratios and claims of a selected global mechanism are not justified.',
            'acceptance': 'A broader in vivo-by-in vitro start matrix or equivalent sensitivity analysis is completed; local failure is diagnosed; conclusions are reported as stable derived-function directions with basin-dependent ranges.'
        },
        {
            'title': 'Practical identifiability and bound sensitivity are not quantified',
            'review': 'Perform likelihood-based or equivalent practical-identifiability analysis for headline parameters and derived biological functions.',
            'manuscript': 'Repository text interprets several individual fitted parameters as biological axes but does not report formal identifiability, likelihood profiles, bound sensitivity, or uncertainty intervals. Figure 4/6 acknowledge multiple parameter regimes but do not resolve parameter determination.',
            'evidence': 'The top-ten separate in vivo objectives span only 14.1193-14.1748 while multiple parameters vary widely; sigma_burden is at its upper bound in all ten. Joint context values also hit bounds, and multiple soft dimensions are saturated. See `analysis/independent_parameter_values_long.csv`, `analysis/independent_parameter_boundary_hits.csv`, and `analysis/joint_soft_coupling_parameter_summary.csv`.',
            'why': 'Near-equivalent objective values across broad parameter ranges mean the data may constrain combinations/functions rather than individual parameters. Bound-limited estimates can create apparent precision and distort context ratios.',
            'external': 'Raue et al. define practical non-identifiability as arising from the amount and quality of data, distinguish it from structural non-identifiability, and propose profile likelihood to detect flat directions and derive likelihood-based intervals. Relevant locations: Abstract; Sections 2.2, 2.3, and 6. Bioinformatics 2009; doi:10.1093/bioinformatics/btp358.',
            'verification': 'Verified diagnostic pattern; formal identifiability analysis has not been supplied.',
            'impact': 'High. Parameter-level biological interpretations and exact fold changes may not be uniquely supported.',
            'acceptance': 'Profiles/bootstrap/envelopes are supplied for headline parameters and derived lambda, p_mis, and survival functions; bounded versus unbounded directions and function-level intervals are reported.'
        },
        {
            'title': 'Most fixed-O2 response curves and all six warm-start representatives lack strong spectral reliability',
            'review': 'Make spectral-gap reliability a primary part of Figure 6 and the joint-selection interpretation.',
            'manuscript': 'Repository Figure 6 and Results emphasize multiple response classes and objective overlap, but the main narrative does not report that only a minority of curves are reliable or that no selected warm start is reliable.',
            'evidence': '`joint_pre.zip` supplies 500 seeds x 201 O2 values. Reliability counts are 141 reliable, 168 caution, and 191 unreliable. Warm starts 25, 311, 322 are caution; 138, 290, 366 are unreliable. See `analysis/fixed_o2_reliability_counts.csv` and `analysis/joint_warm_start_fixed_o2_crosswalk.csv`.',
            'why': 'The dominant eigenvector is an asymptotic descriptor. A small first-to-second eigenvalue gap implies slow separation and sensitivity of the claimed long-run composition; raw shape labels can then overstate mechanistic stability.',
            'external': 'The reliability concern follows directly from the model’s reported spectral-gap definition; no outside biological premise is required.',
            'verification': 'Verified from all 100,500 supplied attractor rows and per-seed classifications.',
            'impact': 'High. It weakens claims that the fitted solution set defines stable long-term oxygen-ploidy regimes and directly affects the representativeness of the joint starts.',
            'acceptance': 'Figure 6 stratifies or annotates reliability; all class counts are cross-tabulated with reliability; warm-start statuses are disclosed; finite-time or threshold sensitivity is provided.'
        },
        {
            'title': 'The Figure 6 objective overlay is not the fitted MAP objective',
            'review': 'Recompute the selection diagnostic using the full MAP objective or label it explicitly as raw burden-plus-ploidy likelihood.',
            'manuscript': 'Repository Figure 6D caption describes objective distributions and lowest-objective fits without distinguishing the objective field selected by the fixed-O2 preprocessing pipeline from the actual separate-fit MAP objective.',
            'evidence': 'For all 500 seeds, the fixed-O2 by-seed table records objective_source = raw_likelihood. The selected value is the raw burden-plus-ploidy likelihood. Its Spearman correlation with separate MAP is 0.429; objective_data versus MAP is 0.915. See `analysis/fixed_o2_objective_definition_audit.csv` and `analysis/fixed_o2_objective_spearman_correlations.csv`.',
            'why': 'Different objective definitions can reorder seeds and classes. A lowest-decile overlay based on one score cannot be described as a MAP-quality subset based on another score.',
            'external': 'No external reference is needed; this is a direct metric-definition mismatch.',
            'verification': 'Verified for all 500 fixed-O2 seed records.',
            'impact': 'High for Figure 6 interpretation. The current overlay cannot select a preferred MAP response class.',
            'acceptance': 'The panel uses an explicitly reconstructed MAP objective or is consistently labeled raw likelihood; the exact formula and seed-level crosswalk are released.'
        },
        {
            'title': 'Necrosis export is broken and the draft equation differs from the implemented objective by a factor of two',
            'review': 'Repair the per-sample necrosis export and use one identical loss definition in code, reports, and manuscript.',
            'manuscript': 'Repository Supplementary Methods defines terminal necrosis loss with a one-half multiplier. The main text treats necrosis as an active calibration stream but does not reveal the export failure.',
            'evidence': 'All inspected `necrosis_fit.tsv` files contain observed values but missing predicted fields. The predicted terminal fraction can be reconstructed from `burden_fit.tsv`. The mean standardized squared clipped-logit residual without 1/2 exactly matches the reported objective (0.631-1.118 across the separate top ten); the half-scaled calculation is exactly half. See `analysis/invivo_top10_fit_quality.csv` and `analysis_tables/invivo_necrosis_reconstruction.csv`.',
            'why': 'A result cannot be independently audited when the standard per-sample export is incomplete. The factor-of-two difference also changes the effective objective weight relative to other terms.',
            'external': 'No external reference is needed; the result follows from exact objective reconstruction.',
            'verification': 'Verified across 10 separate in vivo and 60 joint selected outputs.',
            'impact': 'High. It affects method accuracy, component weighting, reporting, and reproducibility.',
            'acceptance': 'Export contains predicted fraction and contribution for every sample; automated reconstruction equals fit_summary; equation and code use the same normalization.'
        },
        {
            'title': 'The 2N in vitro trajectory and individual in vivo fit are overstated',
            'review': 'Revise trajectory language to match the mixed distribution and disclose the lineage/cohort-level nature of the predictions.',
            'manuscript': 'Repository in vitro Results describe a transient high-ploidy expansion followed by lower-ploidy reshaping. Repository in vivo language can be read as fitting individual tumor trajectories, although the model shares predictions within cohort/harvest structure.',
            'evidence': 'Late 2N deprived observations have means 66.85 and 88.05, but selected predictions are shared at 63.72-64.40. In vivo tumors with the same starting cohort and harvest day receive the same model trajectory because harvest_init_multiplier is false and no mouse-specific latent environment is fitted. Absolute fit-quality metrics include log burden RMSE 0.663-0.680 and terminal distribution TV 0.681-0.714.',
            'why': 'A model can generate a high-chromosome component without showing an overall return of the population mean. Shared cohort predictions cannot explain individual heterogeneity even if the total objective is acceptable.',
            'external': 'No external premise is required; these are direct model-structure and prediction-table observations.',
            'verification': 'Verified across selected in vitro/in vivo outputs and config.',
            'impact': 'Medium-High. It affects the scope of biological claims and the perceived quality of fit.',
            'acceptance': 'Results distinguish distribution components from mean trajectory, show late branch residuals, report absolute fit metrics, and define the model as cohort/lineage-level.'
        },
        {
            'title': 'CIN-associated loss conflates biological nonviability with numerical boundary routing',
            'review': 'Separate the components and perform chromosome-grid sensitivity before interpreting the magnitude biologically.',
            'manuscript': 'Repository model definitions route nonviable daughters and out-of-grid daughter/WGD mass to the same CIN-associated dead pool. The in vitro narrative uses this pool to support altered-daughter filtering.',
            'evidence': 'The retained grid is N=22-154. Model code and manuscript equations route out-of-domain mass to CIN-dead. Selected in vitro fits show max CIN-dead 32.15-33.40% versus direct hypoxia-dead 1.74-1.82%, making component semantics consequential.',
            'why': 'Numerical truncation can inflate a compartment that is subsequently interpreted as biological lethality. The effect may be strongest near high chromosome states and WGD boundaries.',
            'external': 'No external reference is required; it is a direct consequence of the finite-state implementation.',
            'verification': 'Verified from model semantics and selected outputs; grid sensitivity has not been supplied.',
            'impact': 'Medium-High. It affects the key mechanistic claim about how high-ploidy populations are filtered.',
            'acceptance': 'Missegregation nonviability and boundary-dropped mass are separately exported and plotted; principal conclusions are stable under expanded grids.'
        },
        {
            'title': 'Figure-level reproducibility and provenance are incomplete',
            'review': 'Release a clean, executable figure-to-result package with all intermediate tables and exact code state.',
            'manuscript': 'The repository manuscript references Figures 4-6 and extensive numerical results but does not provide a complete Data/Code Availability and figure-source map. The draft includes unresolved placeholders for dataset/provenance citations.',
            'evidence': 'Selected in vivo runs record commit 54c5d500... with dirty worktrees; selected joint runs record 3d498599... with dirty worktrees; selected in vitro seed directories lack equivalent provenance files. Full fixed-O2/classification tables are available in joint_pre, but no definitive seed-level Figure 4 AUC table was found.',
            'why': 'For a code-sensitive mechanistic model, a commit alone cannot reproduce a dirty run, and a figure without its intermediate table cannot be independently checked.',
            'external': 'Sandve et al. Rules 1, 4, 8, and 10 require tracking the workflow behind every result, versioning custom code, retaining detail underlying summaries, and providing access to scripts/runs/results (PLOS Comput Biol 2013; doi:10.1371/journal.pcbi.1003285). Wilkinson et al. state that FAIR applies to algorithms, tools, and workflows and specify detailed provenance under R1.2 (Scientific Data 2016; Box 2; doi:10.1038/sdata.2016.18).',
            'verification': 'Verified from supplied provenance files and archive inventory.',
            'impact': 'High for publication auditability and long-term reuse.',
            'acceptance': 'Clean tag or archived patch, resolved config, parameter tables, environment/container, all seed IDs, checksums, full intermediate tables including AUC, and an executable figure-generation manifest are released.'
        },
    ]

    for idx, e in enumerate(entries, 1):
        doc.add_heading(f'Major Comment {idx}: {e["title"]}', level=1)
        add_label_paragraph(doc, 'Review comment: ', e['review'])
        add_label_paragraph(doc, 'Manuscript basis: ', e['manuscript'])
        add_label_paragraph(doc, 'Result/code evidence: ', e['evidence'])
        add_label_paragraph(doc, 'Why this is a concern: ', e['why'])
        add_label_paragraph(doc, 'External methodological basis: ', e['external'])
        add_label_paragraph(doc, 'Verification level: ', e['verification'])
        add_label_paragraph(doc, 'Impact: ', e['impact'])
        add_label_paragraph(doc, 'Acceptance criterion: ', e['acceptance'])

    doc.add_heading('Methodological References', level=1)
    refs = [
        'Raue A, Kreutz C, Maiwald T, et al. Structural and practical identifiability analysis of partially observed dynamical models by exploiting the profile likelihood. Bioinformatics. 2009;25(15):1923-1929. doi:10.1093/bioinformatics/btp358. Specific support used here: Abstract; Sections 2.2, 2.3, and 6.',
        'Wattenberg M, Viégas F, Johnson I. How to Use t-SNE Effectively. Distill. 2016. doi:10.23915/distill.00002. Specific support used here: Sections 1, 2, and 3 on hyperparameter dependence, cluster size, and inter-cluster distance.',
        'Sandve GK, Nekrutenko A, Taylor J, Hovig E. Ten Simple Rules for Reproducible Computational Research. PLoS Comput Biol. 2013;9(10):e1003285. doi:10.1371/journal.pcbi.1003285. Specific support used here: Rules 1, 4, 8, and 10.',
        'Wilkinson MD, Dumontier M, Aalbersberg IJJ, et al. The FAIR Guiding Principles for scientific data management and stewardship. Sci Data. 2016;3:160018. doi:10.1038/sdata.2016.18. Specific support used here: applicability to algorithms/tools/workflows and Box 2, principle R1.2.'
    ]
    for ref in refs:
        doc.add_paragraph(ref)

    path = OUT / '06_major_comment_crosswalk.docx'
    doc.save(path)
    return path


if __name__ == '__main__':
    parser = argparse.ArgumentParser(description='Regenerate the two plain-structure Word review documents.')
    parser.add_argument(
        '--out-dir',
        type=Path,
        default=Path(__file__).resolve().parents[1],
        help='Destination directory; defaults to the delivery-package root.',
    )
    args = parser.parse_args()
    OUT = args.out_dir.expanduser().resolve()
    OUT.mkdir(parents=True, exist_ok=True)
    print(build_report())
    print(build_crosswalk())
